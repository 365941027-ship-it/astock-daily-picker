"""报告生成：Markdown 与 Word(.docx)，格式对齐历史日报。"""

from __future__ import annotations

import os
from datetime import date, datetime
from typing import Dict, List

from .config import Config
from .screening import market_stats

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    HAS_DOCX = True
except ImportError:  # pragma: no cover - 未安装 python-docx 时仅输出 md
    HAS_DOCX = False


def _fmt_pct(v) -> str:
    return "—" if v is None else f"{v:+.2f}%"


def _fmt_num(v, digits: int = 2) -> str:
    return "—" if v is None else f"{v:.{digits}f}"


def _yi(v) -> str:
    return "—" if v is None or v == 0 else f"{v / 1e8:.1f}亿"


def build_content(
    result: Dict,
    index_data: Dict[str, Dict],
    cfg: Config,
    data_date: date,
    run_date: date,
    observe_date: date,
    data_ok: bool,
    intraday: bool,
    source_note: str,
    market_verdict: str = "",
) -> Dict:
    """把选股结果组织成结构化内容，供 md/docx 两种渲染器共用。"""
    stats = market_stats(result["snapshot"], cfg) if result.get("snapshot") else {}
    sections: List[Dict] = []

    # ---- 市场环境 ----
    market_rows = []
    if index_data:
        for name, bar in index_data.items():
            market_rows.append([
                name,
                _fmt_num(bar.get("close")),
                _fmt_pct(bar.get("pct_chg")),
                _yi(bar.get("amount")),
            ])
    if market_rows:
        sections.append({
            "kind": "table",
            "headers": ["指数", "收盘", "涨跌幅", "成交额"],
            "rows": market_rows,
        })
    if stats:
        sections.append({
            "kind": "para",
            "text": (
                f"全市场 {stats['total']} 只：上涨 {stats['up']}，下跌 {stats['down']}，"
                f"平盘 {stats['flat']}；涨停（按各板块涨跌幅限制口径）{stats['limit_up']} 只。"
            ),
        })

    # ---- 筛选逻辑 ----
    sections.append({"kind": "h2", "text": "筛选逻辑"})
    sections.append({
        "kind": "bullet",
        "items": [
            "非 ST，剔除新股/退市风险标的（上市不足 60 个交易日）。",
            "今日收红。",
            f"成交额大于 {cfg.min_amount / 1e8:.0f} 亿元，换手率大于 {cfg.min_turnover:.0f}%。",
            "收盘站上 20 日线，优先 5 日线大于 10 日线。",
            "MACD 黄白线在 0 轴上方，或 MACD 金叉/柱体改善。",
            f"KDJ 金叉或低中位拐头，但不过热（J ≤ {cfg.kdj_max_j:.0f}）。",
            f"近 5 日涨幅 ≤ {cfg.max_5d_gain:.0f}%、10 日涨幅 ≤ {cfg.max_10d_gain:.0f}%，"
            "涨停或接近涨停的票只列为“等回踩”。",
        ],
    })

    # ---- 今日结论 ----
    if market_verdict == "不适合入场":
        sections.append({"kind": "h1", "text": "今日结论：大盘不宜入场，仅做空仓观察"})
        sections.append({
            "kind": "para",
            "text": (
                f"缠论大盘研判判定：{market_verdict}。"
                "多数指数位于缠论中枢下方，市场环境偏弱，按纪律只做空仓观察，"
                "本轮不推荐任何个股。"
            ),
        })
        sections.append({
            "kind": "para",
            "text": f"操作结论：{observe_date.isoformat()} 空仓观察，等待指数站稳中枢或出现明确转强信号。",
        })
    elif not data_ok:
        sections.append({"kind": "h1", "text": "今日结论：不硬筛个股，只做空仓观察"})
        sections.append({
            "kind": "para",
            "text": (
                f"{data_date.isoformat()} 无法获取公开行情数据（{source_note}）。"
                "没有当日收盘价、成交额、换手率、均线、MACD/KDJ 和板块强弱的交叉核验，"
                "就不能硬说某只票明天能看。"
            ),
        })
        sections.append({
            "kind": "para",
            "text": (
                f"操作结论：{observe_date.isoformat()} 先空仓观察。"
                "旧观察池里的股票只能当作复盘样本，不能直接当作新候选。"
            ),
        })
    elif intraday:
        sections.append({"kind": "h1", "text": "今日结论：盘中数据，仅作预览，不列正式候选"})
        sections.append({
            "kind": "para",
            "text": (
                f"运行时间为 {datetime.now().astimezone().strftime('%H:%M')}，"
                "未到当日 15:05 收盘复核，当前为盘中数据。"
                "短线候选必须用当日收盘后的数据确认，因此本轮只输出预览统计，不列正式个股。"
            ),
        })
        sections.append({
            "kind": "para",
            "text": (
                f"如需盘中预览候选，可加 --allow-intraday 重跑；正式候选请等收盘后运行 "
                f"python pick_stocks.py --date {data_date.isoformat()}。"
            ),
        })
    elif not result["priority"]:
        sections.append({"kind": "h1", "text": "今日结论：没有通过全部筛选条件的个股，空仓观察"})
        sections.append({
            "kind": "para",
            "text": (
                f"{data_date.isoformat()} 收盘后，全市场没有同时满足“收红、流动性、站上20日线、"
                "MACD 改善、KDJ 拐头且不过热、涨幅不透支”的个股。"
                "按纪律处理：不硬筛、不补编、不沿用旧票。"
            ),
        })

    if data_ok and not intraday and market_verdict != "不适合入场":
        # ---- 优先观察 ----
        sections.append({"kind": "h1", "text": f"优先观察（{len(result['priority'])} 只）"})
        if result["priority"]:
            rows = [[
                c.code, c.name,
                _fmt_num(c.close),
                _fmt_pct(c.pct_chg),
                _fmt_num(c.turnover) + "%",
                c.structure,
                c.note,
            ] for c in result["priority"]]
            sections.append({
                "kind": "table",
                "headers": ["代码", "名称", "收盘", "涨幅", "换手", "技术结构", "观察要点"],
                "rows": rows,
            })
        else:
            sections.append({"kind": "para", "text": "本轮没有个股进入优先观察。"})

        # ---- 强势但不宜追高 ----
        sections.append({"kind": "h1", "text": "强势但不宜追高"})
        if result["strong"]:
            rows = [[
                c.code, c.name,
                _fmt_num(c.close),
                _fmt_pct(c.pct_chg),
                c.note,
            ] for c in result["strong"]]
            sections.append({
                "kind": "table",
                "headers": ["代码", "名称", "收盘", "涨幅", "风险点 / 处理方式"],
                "rows": rows,
            })
        else:
            sections.append({"kind": "para", "text": "本轮没有“强势但不宜追高”的标的。"})

        # ---- 剔除或降级观察 ----
        sections.append({"kind": "h1", "text": "剔除或降级观察"})
        if result["excluded"]:
            items = [
                f"{c.code} {c.name}：{'；'.join(c.reasons[:3])}"
                for c in result["excluded"]
            ]
            sections.append({"kind": "bullet", "items": items})
        else:
            sections.append({"kind": "para", "text": "本轮没有需要单独列出的剔除样本。"})

    # ---- 次日入场触发条件 ----
    sections.append({"kind": "h1", "text": "次日入场触发条件"})
    sections.append({
        "kind": "numbered",
        "items": [
            "回踩 5 日线或 10 日线不破，分时出现底背驰或放量重新站回均价线。",
            "回踩前一日实体中部或关键均线后缩量止跌，随后放量突破分时平台。",
            "突破今日高点但不爆量，且板块同步走强。",
        ],
    })

    # ---- 失效条件 ----
    sections.append({"kind": "h1", "text": "失效条件"})
    sections.append({
        "kind": "bullet",
        "items": [
            "跌破 10 日线且无法快速收回。",
            "MACD 黄白线重新走弱并靠近/跌破 0 轴。",
            "高开冲高后放量回落，跌回前一日实体下半区。",
            "所属板块由强转弱，个股独强但量能失控。",
        ],
    })

    # ---- 数据与风险说明 ----
    sections.append({"kind": "h1", "text": "数据与风险说明"})
    sections.append({
        "kind": "bullet",
        "items": [
            f"数据源：{source_note}；数据日期：{data_date.isoformat()}；"
            f"运行时间：{run_date.isoformat()}。",
            "本报告只做短线技术观察，不构成投资建议，不承诺收益。",
            "缠论口径里，二买、三买都要看盘中“回踩不破、拉起有效”，需结合分时和板块确认。",
            "KDJ 只辅助节奏，不能单独作为买点；所有信号以次日盘中实际走势为准。",
        ],
    })

    return {
        "title": f"A股短线候选观察清单 - {data_date.isoformat()}",
        "intro": (
            f"数据口径：{source_note}。报送日期 {run_date.isoformat()}，"
            f"观察日 {observe_date.isoformat()}（下一交易日）。"
            "以下仅为技术观察池，不构成投资建议。"
        ),
        "sections": sections,
    }


def render_markdown(content: Dict) -> str:
    lines = [f"# {content['title']}", "", f"> {content['intro']}", ""]
    for sec in content["sections"]:
        kind = sec["kind"]
        if kind == "h1":
            lines += [f"## {sec['text']}", ""]
        elif kind == "h2":
            lines += [f"### {sec['text']}", ""]
        elif kind == "para":
            lines += [sec["text"], ""]
        elif kind == "bullet":
            lines += [f"- {b}" for b in sec["items"]]
            lines.append("")
        elif kind == "numbered":
            lines += [f"{i}. {b}" for i, b in enumerate(sec["items"], 1)]
            lines.append("")
        elif kind == "table":
            headers = sec["headers"]
            rows = sec["rows"]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("|" + "|".join(["---"] * len(headers)) + "|")
            for r in rows:
                cells = [str(c).replace("|", "／") for c in r]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def render_docx(content: Dict, path: str) -> None:
    """按历史日报样式生成 Word 报告（依赖 python-docx）。"""
    if not HAS_DOCX:
        raise RuntimeError("未安装 python-docx，无法生成 .docx，请安装后重试")

    def set_east_asia_font(run, name="Microsoft YaHei"):
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)

    def add_run(p, text, bold=False, color=None, size=None):
        run = p.add_run(text)
        set_east_asia_font(run)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        if size:
            run.font.size = Pt(size)
        return run

    def shade(cell, fill):
        pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        pr.append(shd)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 12.5, RGBColor(46, 116, 181)),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, content["title"], bold=True, color=(31, 77, 120), size=18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub, content["intro"], color=(90, 90, 90), size=9)

    for sec in content["sections"]:
        kind = sec["kind"]
        if kind == "h1":
            doc.add_heading(sec["text"], level=1)
        elif kind == "h2":
            doc.add_heading(sec["text"], level=2)
        elif kind == "para":
            p = doc.add_paragraph()
            add_run(p, sec["text"])
        elif kind in ("bullet", "numbered"):
            style = "List Bullet" if kind == "bullet" else "List Number"
            for item in sec["items"]:
                p = doc.add_paragraph(style=style)
                add_run(p, item)
        elif kind == "table":
            headers = sec["headers"]
            rows = sec["rows"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr[i].text = ""
                p = hdr[i].paragraphs[0]
                add_run(p, h, bold=True)
                shade(hdr[i], "F2F4F7")
            for r in rows:
                cells = table.add_row().cells
                for i, v in enumerate(r):
                    cells[i].text = ""
                    add_run(cells[i].paragraphs[0], str(v))
            doc.add_paragraph()

    doc.save(path)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "A股短线候选_2026-07-09.docx"


def set_east_asia_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_run(paragraph, text, bold=False, color=None):
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_para(doc, text="", style=None, bold=False):
    p = doc.add_paragraph(style=style)
    if text:
        add_run(p, text, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    add_run(p, text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    add_run(p, text)
    return p


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color in [
    ("Heading 1", 16, RGBColor(46, 116, 181)),
    ("Heading 2", 13, RGBColor(46, 116, 181)),
    ("Heading 3", 12, RGBColor(31, 77, 120)),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(12 if name != "Heading 1" else 16)
    style.paragraph_format.space_after = Pt(6)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("A股次日短线观察日报")
set_east_asia_font(run)
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(31, 77, 120)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(sub, "报送日期：2026年7月9日  |  观察日：2026年7月10日")

p = add_para(doc)
add_run(p, "下面按", bold=False)
add_run(p, "明天 2026年7月10日", bold=True)
add_run(p, "怎么盯盘来讲。先记住一句话：")

p = add_para(doc)
add_run(p, "这些票不是明天开盘就买，而是观察它们能不能“跌不破、拉得起、站得住”。", bold=True)

add_para(doc, "如果看不懂复杂指标，就只看三件事：")
add_number(doc, "别追高：开盘一下涨很多，不急着买。")
add_number(doc, "看支撑：跌到关键价位附近，能不能跌不动。")
add_number(doc, "看反拉：跌不动以后，能不能重新往上拉。")

h = add_para(doc, "今日结论：不硬筛个股，只做空仓观察", style="Heading 1")
p = add_para(doc)
add_run(p, "今天按日历是周四，且不属于常见法定节假日区间，正常情况下应属于 A 股交易日。")
add_run(p, "但本次运行环境无法访问公开行情接口：Python 侧访问东方财富接口出现 DNS 解析失败，Node 侧访问同一公开接口返回 fetch failed；普通网页搜索也没有拿到可复核的 2026年7月9日收盘行情页。", bold=True)

p = add_para(doc)
add_run(p, "因此，今天不把任何股票列为新的“次日短线可观察入场个股”。")
add_run(p, "原因很简单：没有当日收盘价、5日线、10日线、成交额、换手率、MACD/KDJ 和板块强弱的交叉核验，就不能硬说某只票明天能看。", bold=True)

p = add_para(doc)
shade_paragraph(p, "F2F4F7")
add_run(p, "操作结论：明天先空仓观察。旧观察池里的股票只能当作复盘样本，不能直接当作 2026年7月10日的新候选。")

add_para(doc, "为什么今天不硬选", style="Heading 1")
add_bullet(doc, "短线候选必须用当日收盘后的数据确认，尤其是收盘价是否站上关键均线、成交额是否足够、MACD 是否改善。")
add_bullet(doc, "缠论口径里，二买、三买都要看“回踩是否不破”和“拉起是否有效”，这必须依赖最新行情。")
add_bullet(doc, "KDJ 只能辅助节奏，不能单独作为买点；没有当日数据时，更不能只凭旧清单延续判断。")
add_bullet(doc, "如果明天开盘前仍拿不到 2026年7月9日的完整日线数据，就继续按空仓观察处理。")

add_para(doc, "历史观察池怎么处理", style="Heading 1")
p = add_para(doc)
add_run(p, "工作区最近一份历史清单是 2026年7月7日的观察池，里面出现过国恩股份、中科三环、硅宝科技、恒星科技、大港股份等。")
add_run(p, "这些名字今天不升级为正式候选。", bold=True)
add_run(p, "如果你明天还想盯，只能按下面的“重新确认规则”看，不能因为它们曾经在旧清单里就追。")

add_para(doc, "旧观察池重新确认规则", style="Heading 2")
add_bullet(doc, "先看开盘后有没有明显弱于板块；如果板块走弱、个股独自冲高，先不动。")
add_bullet(doc, "看它有没有回踩到 5日线或 10日线附近后跌不破，而不是一开盘就追高。")
add_bullet(doc, "看分时能不能重新站回当天均价线；站不回去，就说明拉不动。")
add_bullet(doc, "看成交量是不是一路放量砸盘；如果是放量下跌，不要把它当成洗盘。")
add_bullet(doc, "跌破关键均线或前一日结构低点后拉不回来，就直接放弃。")

add_para(doc, "明天怎么看", style="Heading 1")
p = add_para(doc)
add_run(p, "明天的重点不是找谁涨得快，而是等市场自己给答案。")
add_run(p, "如果盘中有股票回踩后不破、重新拉回均价线、成交量没有失控，再考虑把它放进临时观察。", bold=True)

add_para(doc, "可以观察入场的信号", style="Heading 2")
add_bullet(doc, "跌到关键均线或前一日平台附近，不继续破。")
add_bullet(doc, "分时图重新往上走，或重新站回当天均价线。")
add_bullet(doc, "成交量不是一路放量砸盘，而是回落缩量、拉起时温和放量。")
add_bullet(doc, "所属板块同步回暖，个股不是孤零零地硬拉。")

add_para(doc, "不要买的情况", style="Heading 2")
add_bullet(doc, "高开后一路往下砸，或冲高后快速回落。")
add_bullet(doc, "跌破关键支撑后拉不回来。")
add_bullet(doc, "成交量放大但价格站不住，说明抛压重。")
add_bullet(doc, "看不清支撑价、看不清退出线，只是因为它涨了就想追。")

p = add_para(doc)
add_run(p, "简单说：", bold=True)
add_run(p, "今天没有经过当日行情复核的正式候选，明天只看“回踩后还能不能站起来”。站不起来就放弃。")

add_para(doc, "最简单的明天操作规则", style="Heading 1")
add_para(doc, "你明天可以这样看：")

p = add_para(doc)
add_run(p, "第一步：开盘不急着买。", bold=True)
add_run(p, " 前30分钟先看方向，尤其不要追高开很多的票。")

p = add_para(doc)
add_run(p, "第二步：只买“回踩后重新拉起”的票。", bold=True)
add_run(p, " 也就是先跌一点，跌到关键支撑附近不破，然后重新往上走。")

p = add_para(doc)
add_run(p, "第三步：只选最强的一两只。", bold=True)
add_run(p, " 不要把候选都买了。明天如果只有一只走势符合，就只看一只；没有符合的，就空着。")

p = add_para(doc)
add_run(p, "第四步：一旦买错，要有退出线。", bold=True)
add_run(p, " 比如买入后跌破你观察的支撑价，并且拉不回来，就不要硬扛。")

add_para(doc, "一句话版：")
p = add_para(doc)
add_run(p, "明天不是看谁涨得最快，而是看谁回落后还能重新站起来。能站起来的才值得看，站不起来的直接放弃。", bold=True)

add_para(doc, "数据缺口说明", style="Heading 1")
add_bullet(doc, "公开行情接口未能从当前环境访问，无法交叉核验指数环境、全市场宽基表现、成交额、行业强弱和个股日线。")
add_bullet(doc, "本报告因此不列正式个股，不承诺收益，也不给确定性买入指令。")
add_bullet(doc, "若补充 2026年7月9日收盘后的个股日线数据，可重新筛出 3-8 只候选并补齐具体支撑观察区间。")

doc.save(OUT)
print(OUT)
